from docx import Document
from docx.shared import Pt
import json
from pydantic import BaseModel, Field

class StructuredOutput(BaseModel):
    """
    Structured output
    """
    answer: str = Field(..., description="La risposta dell'LLM alla domanda.")
    confidence: float = Field(..., description="Confidenza dell'LLM nella risposta, da 0 a 1.")
    confidence_explanation: str = Field(..., description="Spiegazione della confidenza dell'LLM nella risposta.")
    sources: str = Field(..., description="Fonti utilizzate dall'LLM per generare la risposta.")

def generate_response(query, llm):
    """    Generate a response from the LLM using the structured output format.
    Args:
        query (str): The input query for the LLM.
        llm (structured_llm): The structured LLM instance.
    Returns:
        StructuredOutput: The structured response from the LLM.
    """
    resp = llm.complete(query)
    str_resp = resp.model_dump_json()
    dict_resp = json.loads(str_resp)
    text = dict_resp["text"]
    text_dict = json.loads(text)
    return text_dict

def create_document(text_dict, filename='llm_response.docx'):
    """
    Create a Word document with the structured response from the LLM.
    """

    doc = Document()
    doc.add_heading('LLM structured Response', level=1)
    doc.add_heading('Answer', level=2)
    doc.add_paragraph(f"{text_dict['answer']}")
    doc.add_heading('LLM Confidence Score', level=2)
    doc.add_paragraph(f"{text_dict['confidence']}")
    doc.add_heading('LLM Confidence Explanation', level=2)
    doc.add_paragraph(f"{text_dict['confidence_explanation']}")
    doc.add_heading('LLM Sources', level=2)
    if text_dict.get('sources'):
        doc.add_paragraph(f"{text_dict['sources']}")
    else:
        doc.add_paragraph("No sources provided.")
    doc.add_heading('Level of Satisfaction with the Answer', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = '0%'
    table.cell(0, 1).text = '25%'
    table.cell(0, 2).text = '50%'
    table.cell(0, 3).text = '75%'
    table.cell(0, 4).text = '100%'
    doc.save(filename)