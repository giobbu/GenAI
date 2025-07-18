from docx import Document
from docx.shared import Pt
import json
from pydantic import BaseModel, Field

class StructuredOutput(BaseModel):
    """
    Structured output
    """
    risposta: str = Field(..., description="La risposta dell'LLM alla domanda.")
    punteggio_confidenza: float = Field(..., description="Confidenza dell'LLM nella risposta, nell'intervallo 0.0 e 1.0.")
    spiegazione_confidenza: str = Field(..., description="Spiegazione della confidenza dell'LLM nella risposta.")
    riferimenti: str = Field(..., description="Riferimenti utilizzati dall'LLM per generare la risposta.")

def generate_response(query, llm):
    """Generate a response from the LLM using the structured output format.
    Args:
        query (str): The input query for the LLM.
        llm (structured_llm): The structured LLM instance.
    Returns:
        StructuredOutput: The structured response from the LLM.
    """
    resp = llm.complete(query)
    str_resp = resp.model_dump_json()
    dict_resp = json.loads(str_resp)
    text = dict_resp["text"]
    text_dict = json.loads(text)
    return text_dict

def create_document(text_dict, filename='llm_response.docx'):
    """
    Create a Word document with the structured response from the LLM.
    """

    doc = Document()
    doc.add_heading('LLM con risposta strutturata', level=1)
    doc.add_heading('Risposta', level=2)
    doc.add_paragraph(f"{text_dict['risposta']}")

    doc.add_heading('Punteggio di confidenza da parte LLM', level=2)
    doc.add_paragraph(f"{text_dict['punteggio_confidenza']}")

    doc.add_heading('Spiegazione della confidenza  da parte LLM', level=2)
    doc.add_paragraph(f"{text_dict['spiegazione_confidenza']}")

    doc.add_heading('Riferimenti LLM', level=2)
    if text_dict.get('riferimenti'):
        doc.add_paragraph(f"{text_dict['riferimenti']}")
    else:
        doc.add_paragraph("Nessuna fonte disponibile.")
    
    doc.add_heading('Livello di soddisfazione da parte dell utilizzatore', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)
    table.cell(0, 0).text = '0%'
    table.cell(0, 1).text = '25%'
    table.cell(0, 2).text = '50%'
    table.cell(0, 3).text = '75%'
    table.cell(0, 4).text = '100%'
    doc.save(filename)